"""
text_extraction.py - Extraccion de texto de documentos para busqueda full-text

Este modulo proporciona funciones para extraer texto de diferentes tipos
de documentos (PDF, Word, Excel, texto plano) para permitir la busqueda
full-text en el contenido de los documentos.

Formatos soportados:
- PDF (.pdf)
- Microsoft Word (.docx)
- Microsoft Excel (.xlsx)
- Texto plano (.txt, .csv, .json, .xml)
- RTF (.rtf)

Autor: Equipo de Desarrollo LegalFlow
"""

import logging
import io
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extrae texto de un archivo PDF.

    Args:
        file_content: Contenido binario del archivo PDF

    Returns:
        str: Texto extraido del PDF
    """
    try:
        from pypdf import PdfReader

        # Crear un objeto BytesIO para leer el PDF
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return '\n'.join(text_parts)

    except Exception as e:
        logger.error(f"Error extrayendo texto de PDF: {str(e)}")
        return ''


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extrae texto de un archivo Microsoft Word (.docx).

    Args:
        file_content: Contenido binario del archivo Word

    Returns:
        str: Texto extraido del documento
    """
    try:
        from docx import Document

        # Crear un objeto BytesIO para leer el documento
        doc_file = io.BytesIO(file_content)
        document = Document(doc_file)

        text_parts = []

        # Extraer texto de parrafos
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extraer texto de tablas
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return '\n'.join(text_parts)

    except Exception as e:
        logger.error(f"Error extrayendo texto de DOCX: {str(e)}")
        return ''


def extract_text_from_xlsx(file_content: bytes) -> str:
    """
    Extrae texto de un archivo Microsoft Excel (.xlsx).

    Args:
        file_content: Contenido binario del archivo Excel

    Returns:
        str: Texto extraido de todas las hojas del documento
    """
    try:
        from openpyxl import load_workbook

        # Crear un objeto BytesIO para leer el archivo
        excel_file = io.BytesIO(file_content)
        workbook = load_workbook(excel_file, data_only=True)

        text_parts = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"--- Hoja: {sheet_name} ---")

            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))

                if row_values:
                    text_parts.append(' | '.join(row_values))

        return '\n'.join(text_parts)

    except Exception as e:
        logger.error(f"Error extrayendo texto de XLSX: {str(e)}")
        return ''


def extract_text_from_plain_text(file_content: bytes, encoding: str = 'utf-8') -> str:
    """
    Extrae texto de un archivo de texto plano.

    Args:
        file_content: Contenido binario del archivo
        encoding: Codificacion del archivo (default: utf-8)

    Returns:
        str: Texto del archivo
    """
    try:
        # Intentar decodificar con la codificacion especificada
        return file_content.decode(encoding)
    except UnicodeDecodeError:
        # Si falla, intentar con latin-1
        try:
            return file_content.decode('latin-1')
        except Exception:
            return ''
    except Exception as e:
        logger.error(f"Error extrayendo texto plano: {str(e)}")
        return ''


def extract_text(file_content: bytes, mime_type: str, filename: str = '') -> str:
    """
    Extrae texto de un archivo basandose en su tipo MIME.

    Esta funcion es el punto de entrada principal para la extraccion
    de texto. Detecta el tipo de archivo y llama a la funcion
    de extraccion apropiada.

    Args:
        file_content: Contenido binario del archivo
        mime_type: Tipo MIME del archivo (ej: application/pdf)
        filename: Nombre del archivo para deteccion por extension

    Returns:
        str: Texto extraido del documento, o cadena vacia si no se puede extraer
    """
    # Normalizar mime_type
    mime_type = (mime_type or '').lower()
    filename = (filename or '').lower()

    # Mapeo de tipos MIME a funciones de extraccion
    mime_extractors = {
        'application/pdf': extract_text_from_pdf,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': extract_text_from_docx,
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': extract_text_from_xlsx,
        'text/plain': extract_text_from_plain_text,
        'text/csv': extract_text_from_plain_text,
        'application/json': extract_text_from_plain_text,
        'application/xml': extract_text_from_plain_text,
        'text/xml': extract_text_from_plain_text,
        'text/html': extract_text_from_plain_text,
        'application/rtf': extract_text_from_plain_text,
    }

    # Buscar extractor por tipo MIME
    extractor = mime_extractors.get(mime_type)

    # Si no se encuentra por MIME, intentar por extension
    if not extractor and filename:
        extension_extractors = {
            '.pdf': extract_text_from_pdf,
            '.docx': extract_text_from_docx,
            '.xlsx': extract_text_from_xlsx,
            '.txt': extract_text_from_plain_text,
            '.csv': extract_text_from_plain_text,
            '.json': extract_text_from_plain_text,
            '.xml': extract_text_from_plain_text,
            '.html': extract_text_from_plain_text,
            '.htm': extract_text_from_plain_text,
            '.rtf': extract_text_from_plain_text,
        }

        for ext, func in extension_extractors.items():
            if filename.endswith(ext):
                extractor = func
                break

    if extractor:
        try:
            text = extractor(file_content)
            # Limpiar el texto extraido
            return clean_extracted_text(text)
        except Exception as e:
            logger.error(f"Error en extraccion de texto para {filename}: {str(e)}")
            return ''

    logger.info(f"No hay extractor disponible para tipo MIME: {mime_type}, archivo: {filename}")
    return ''


def clean_extracted_text(text: str) -> str:
    """
    Limpia el texto extraido eliminando caracteres innecesarios.

    Args:
        text: Texto a limpiar

    Returns:
        str: Texto limpio
    """
    if not text:
        return ''

    # Normalizar saltos de linea
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Eliminar multiples espacios en blanco consecutivos
    import re
    text = re.sub(r'[ \t]+', ' ', text)

    # Eliminar lineas vacias multiples
    text = re.sub(r'\n\s*\n+', '\n\n', text)

    # Eliminar espacios al inicio y final de cada linea
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    # Limitar a un tamano maximo razonable (500KB de texto)
    max_length = 500 * 1024
    if len(text) > max_length:
        text = text[:max_length]
        logger.warning(f"Texto truncado a {max_length} caracteres")

    return text.strip()


def get_text_preview(text: str, max_length: int = 500) -> str:
    """
    Obtiene una vista previa del texto extraido.

    Args:
        text: Texto completo
        max_length: Longitud maxima de la vista previa

    Returns:
        str: Vista previa del texto
    """
    if not text:
        return ''

    if len(text) <= max_length:
        return text

    # Truncar y agregar indicador de que hay mas texto
    return text[:max_length].rsplit(' ', 1)[0] + '...'
